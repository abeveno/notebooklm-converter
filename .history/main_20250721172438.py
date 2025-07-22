# Tên file: epub_all_in_one_converter.py
# Mô tả: Một ứng dụng GUI All-in-one để chuyển đổi file EPUB sang nhiều định dạng khác nhau.
# Hỗ trợ: PDF, TXT, Markdown, HTML, MOBI/AZW3, DOCX
#
# Hướng dẫn cài đặt các thư viện cần thiết:
# Mở Command Prompt (cmd) hoặc PowerShell và chạy lệnh sau:
# pip install EbookLib beautifulsoup4 lxml xhtml2pdf python-docx markdown

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches
import markdown
import os
import base64
import threading
import re
import subprocess
import sys

class EbookAllInOneConverterApp:
    def __init__(self, root):
        """
        Khởi tạo giao diện người dùng cho ứng dụng All-in-one converter.
        """
        self.root = root
        self.root.title("EPUB All-in-One Converter")
        self.root.geometry("700x500")
        self.root.configure(bg="#f0f0f0")

        self.epub_path = None
        self.output_format = tk.StringVar(value="PDF")
        self.quality_var = tk.StringVar(value="normal")

        # --- Frame chính ---
        main_frame = tk.Frame(self.root, padx=20, pady=20, bg="#f0f0f0")
        main_frame.pack(expand=True, fill=tk.BOTH)

        # --- Tiêu đề ---
        title_label = tk.Label(main_frame, text="EPUB All-in-One Converter", font=("Helvetica", 20, "bold"), bg="#f0f0f0", fg="#333")
        title_label.pack(pady=(0, 10))
        
        subtitle_label = tk.Label(main_frame, text="Chuyển đổi EPUB sang PDF, TXT, Markdown, HTML, MOBI/AZW3, DOCX", font=("Helvetica", 10), bg="#f0f0f0", fg="#666")
        subtitle_label.pack(pady=(0, 20))

        # --- Frame cho việc chọn file ---
        file_frame = tk.Frame(main_frame, bg="#f0f0f0")
        file_frame.pack(fill=tk.X, pady=10)

        # --- Nút chọn file ---
        select_button = tk.Button(file_frame, text="1. Chọn file EPUB", command=self.select_epub_file, font=("Helvetica", 12), bg="#007bff", fg="white", relief=tk.FLAT, padx=15, pady=8)
        select_button.pack(side=tk.LEFT)

        # --- Nhãn hiển thị file đã chọn ---
        self.file_label = tk.Label(file_frame, text="Chưa có file nào được chọn", font=("Helvetica", 10, "italic"), bg="#f0f0f0", fg="#555")
        self.file_label.pack(side=tk.LEFT, padx=(15, 0))

        # --- Frame cho việc chọn định dạng ---
        format_frame = tk.LabelFrame(main_frame, text="2. Chọn định dạng đầu ra", font=("Helvetica", 12, "bold"), bg="#f0f0f0", fg="#333", padx=10, pady=10)
        format_frame.pack(fill=tk.X, pady=20)

        # --- Radio buttons cho các định dạng ---
        formats = [
            ("PDF", "PDF", "Portable Document Format - Định dạng tài liệu di động"),
            ("TXT", "TXT", "Plain Text - Văn bản thuần túy"),
            ("Markdown", "MD", "Markdown - Định dạng văn bản có cú pháp đánh dấu"),
            ("HTML", "HTML", "HyperText Markup Language - Ngôn ngữ đánh dấu siêu văn bản"),
            ("DOCX", "DOCX", "Microsoft Word Document - Tài liệu Word"),
            ("MOBI/AZW3", "MOBI", "Amazon Kindle Format - Định dạng sách điện tử Kindle")
        ]

        # Tạo grid layout cho radio buttons
        for i, (display_name, value, description) in enumerate(formats):
            row = i // 2
            col = i % 2
            
            frame = tk.Frame(format_frame, bg="#f0f0f0")
            frame.grid(row=row, column=col, sticky="w", padx=10, pady=5)
            
            radio = tk.Radiobutton(frame, text=display_name, variable=self.output_format, value=value, 
                                 font=("Helvetica", 11, "bold"), bg="#f0f0f0", fg="#333")
            radio.pack(anchor="w")
            
            desc_label = tk.Label(frame, text=description, font=("Helvetica", 8), bg="#f0f0f0", fg="#666")
            desc_label.pack(anchor="w", padx=(20, 0))

        # --- Advanced Options Frame ---
        advanced_frame = tk.LabelFrame(main_frame, text="Tùy chọn nâng cao", font=("Helvetica", 10), bg="#f0f0f0", fg="#666", padx=10, pady=5)
        advanced_frame.pack(fill=tk.X, pady=10)
        
        # Quality options
        self.quality_var = tk.StringVar(value="normal")
        quality_frame = tk.Frame(advanced_frame, bg="#f0f0f0")
        quality_frame.pack(anchor="w")
        
        tk.Label(quality_frame, text="Chất lượng:", font=("Helvetica", 9), bg="#f0f0f0").pack(side=tk.LEFT)
        tk.Radiobutton(quality_frame, text="Nhanh", variable=self.quality_var, value="fast", bg="#f0f0f0", font=("Helvetica", 8)).pack(side=tk.LEFT, padx=5)
        tk.Radiobutton(quality_frame, text="Bình thường", variable=self.quality_var, value="normal", bg="#f0f0f0", font=("Helvetica", 8)).pack(side=tk.LEFT, padx=5)
        tk.Radiobutton(quality_frame, text="Cao", variable=self.quality_var, value="high", bg="#f0f0f0", font=("Helvetica", 8)).pack(side=tk.LEFT, padx=5)

        # --- Nút chuyển đổi ---
        self.convert_button = tk.Button(main_frame, text="3. Chuyển đổi", command=self.start_conversion_thread, 
                                      font=("Helvetica", 14, "bold"), bg="#28a745", fg="white", relief=tk.FLAT, 
                                      padx=20, pady=10, state=tk.DISABLED)
        self.convert_button.pack(pady=30)
        
        # --- Progress bar ---
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.pack(fill=tk.X, pady=10)
        
        # --- Vùng hiển thị trạng thái ---
        self.status_label = tk.Label(main_frame, text="Sẵn sàng - Chọn file EPUB để bắt đầu", font=("Helvetica", 11), bg="#f0f0f0", fg="#333")
        self.status_label.pack(pady=10, side=tk.BOTTOM, fill=tk.X)

    def select_epub_file(self):
        """
        Mở hộp thoại để người dùng chọn một file .epub.
        """
        self.epub_path = filedialog.askopenfilename(
            title="Chọn một file EPUB",
            filetypes=[("EPUB files", "*.epub")]
        )
        if self.epub_path:
            # Lấy tên file để hiển thị
            filename = os.path.basename(self.epub_path)
            self.file_label.config(text=f"Đã chọn: {filename}", font=("Helvetica", 10, "normal"))
            self.convert_button.config(state=tk.NORMAL) # Kích hoạt nút chuyển đổi
            self.status_label.config(text="Đã chọn file. Sẵn sàng để chuyển đổi.")
        else:
            self.file_label.config(text="Chưa có file nào được chọn", font=("Helvetica", 10, "italic"))
            self.convert_button.config(state=tk.DISABLED)

    def start_conversion_thread(self):
        """
        Bắt đầu quá trình chuyển đổi trong một luồng riêng để không làm treo giao diện.
        """
        if not self.epub_path:
            messagebox.showerror("Lỗi", "Vui lòng chọn một file EPUB trước.")
            return

        # Lấy định dạng đầu ra
        format_choice = self.output_format.get()
        
        # Định nghĩa extension tương ứng
        extensions = {
            "PDF": ".pdf",
            "TXT": ".txt", 
            "MD": ".md",
            "HTML": ".html",
            "DOCX": ".docx",
            "MOBI": ".mobi"
        }
        
        # Định nghĩa file types cho dialog
        file_types = {
            "PDF": [("PDF files", "*.pdf")],
            "TXT": [("Text files", "*.txt")],
            "MD": [("Markdown files", "*.md")],
            "HTML": [("HTML files", "*.html")],
            "DOCX": [("Word documents", "*.docx")],
            "MOBI": [("MOBI files", "*.mobi")]
        }

        output_path = filedialog.asksaveasfilename(
            title=f"Lưu file {format_choice}",
            defaultextension=extensions[format_choice],
            filetypes=file_types[format_choice]
        )

        if not output_path:
            self.status_label.config(text="Hủy bỏ lưu file.")
            return
            
        # Vô hiệu hóa các nút và bắt đầu progress bar
        self.convert_button.config(state=tk.DISABLED)
        self.progress.start()
        
        self.status_label.config(text=f"Đang chuyển đổi sang {format_choice}, vui lòng chờ...")
        self.root.update_idletasks()

        # Chạy chuyển đổi trong một luồng khác
        conversion_thread = threading.Thread(
            target=self.convert_epub,
            args=(self.epub_path, output_path, format_choice)
        )
        conversion_thread.start()

    def convert_epub(self, epub_path, output_path, output_format):
        """
        Hàm chính thực hiện việc chuyển đổi EPUB sang định dạng được chọn.
        """
        try:
            book = epub.read_epub(epub_path)
            
            if output_format == "PDF":
                self.convert_to_pdf(book, output_path)
            elif output_format == "TXT":
                self.convert_to_txt(book, output_path)
            elif output_format == "MD":
                self.convert_to_markdown(book, output_path)
            elif output_format == "HTML":
                self.convert_to_html(book, output_path)
            elif output_format == "DOCX":
                self.convert_to_docx(book, output_path)
            elif output_format == "MOBI":
                self.convert_to_mobi(epub_path, output_path)
            
            # Cập nhật giao diện khi thành công
            self.root.after(0, self.conversion_successful, output_path, output_format)

        except Exception as e:
            # Cập nhật giao diện khi có lỗi
            self.root.after(0, self.conversion_failed, e)

    def extract_text_content(self, book):
        """
        Trích xuất nội dung văn bản từ EPUB.
        """
        text_content = ""
        
        for item_id in book.spine:
            item = book.get_item_with_id(item_id[0])
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'lxml')
                # Loại bỏ các thẻ script và style
                for script in soup(["script", "style"]):
                    script.decompose()
                text_content += soup.get_text() + "\n\n"
        
        return text_content

    def extract_html_content(self, book):
        """
        Trích xuất nội dung HTML từ EPUB với images embedded.
        """
        # Lấy tất cả các mục từ sách
        items = list(book.get_items())
        
        # Tạo từ điển cho hình ảnh base64
        images_base64 = {}
        for item in items:
            if item.get_type() == ebooklib.ITEM_IMAGE:
                content = item.get_content()
                encoded_image = base64.b64encode(content).decode('utf-8')
                media_type = item.get_media_type()
                images_base64[item.get_name()] = f"data:{media_type};base64,{encoded_image}"

        # Lấy CSS
        css_content = ""
        for item in items:
            if item.get_type() == ebooklib.ITEM_STYLE:
                css_content += item.get_content().decode('utf-8', 'ignore')

        # Xây dựng HTML
        full_html_content = f"<html><head><style>{css_content}</style></head><body>"

        for item_id in book.spine:
            item = book.get_item_with_id(item_id[0])
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'lxml')
                
                # Thay thế đường dẫn hình ảnh
                for img_tag in soup.find_all('img'):
                    src = img_tag.get('src')
                    if src:
                        img_path = os.path.normpath(os.path.join(os.path.dirname(item.get_name()), src))
                        img_key = img_path.replace('\\', '/')
                        if img_key in images_base64:
                            img_tag['src'] = images_base64[img_key]

                full_html_content += str(soup)
        
        full_html_content += "</body></html>"
        return full_html_content

    def convert_to_pdf(self, book, output_path):
        """
        Chuyển đổi EPUB sang PDF.
        """
        full_html_content = self.extract_html_content(book)
        
        with open(output_path, "w+b") as pdf_file:
            pisa_status = pisa.CreatePDF(full_html_content, dest=pdf_file, encoding='UTF-8')

        if pisa_status.err:
            raise Exception(f"Lỗi khi tạo PDF: {pisa_status.err}")

    def convert_to_txt(self, book, output_path):
        """
        Chuyển đổi EPUB sang TXT.
        """
        text_content = self.extract_text_content(book)
        
        with open(output_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(text_content)

    def convert_to_markdown(self, book, output_path):
        """
        Chuyển đổi EPUB sang Markdown.
        """
        markdown_content = ""
        
        for item_id in book.spine:
            item = book.get_item_with_id(item_id[0])
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'lxml')
                
                # Chuyển đổi các thẻ HTML cơ bản sang Markdown
                # Tiêu đề
                for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    level = int(h.name[1])
                    h.string = '#' * level + ' ' + h.get_text() + '\n\n'
                
                # In đậm và in nghiêng
                for strong in soup.find_all(['strong', 'b']):
                    strong.string = '**' + strong.get_text() + '**'
                
                for em in soup.find_all(['em', 'i']):
                    em.string = '*' + em.get_text() + '*'
                
                # Đoạn văn
                for p in soup.find_all('p'):
                    text = p.get_text()
                    if text.strip():
                        p.string = text + '\n\n'
                
                markdown_content += soup.get_text() + "\n\n"
        
        with open(output_path, 'w', encoding='utf-8') as md_file:
            md_file.write(markdown_content)

    def convert_to_html(self, book, output_path):
        """
        Chuyển đổi EPUB sang HTML.
        """
        html_content = self.extract_html_content(book)
        
        with open(output_path, 'w', encoding='utf-8') as html_file:
            html_file.write(html_content)

    def convert_to_docx(self, book, output_path):
        """
        Chuyển đổi EPUB sang DOCX.
        """
        doc = Document()
        
        for item_id in book.spine:
            item = book.get_item_with_id(item_id[0])
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'lxml')
                
                # Xử lý các element
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
                    text = element.get_text().strip()
                    if text:
                        if element.name.startswith('h'):
                            # Tiêu đề
                            level = int(element.name[1])
                            heading = doc.add_heading(text, level=level)
                        else:
                            # Đoạn văn
                            para = doc.add_paragraph(text)
        
        doc.save(output_path)

    def convert_to_mobi(self, epub_path, output_path):
        """
        Chuyển đổi EPUB sang MOBI/AZW3 sử dụng Calibre.
        """
        try:
            # Kiểm tra xem calibre có được cài đặt không
            result = subprocess.run(['ebook-convert', '--version'], 
                                  capture_output=True, text=True, shell=True)
            if result.returncode != 0:
                raise Exception("Calibre chưa được cài đặt. Vui lòng cài đặt Calibre để chuyển đổi sang MOBI.")
            
            # Chạy lệnh chuyển đổi
            cmd = ['ebook-convert', epub_path, output_path]
            result = subprocess.run(cmd, capture_output=True, text=True, shell=True)
            
            if result.returncode != 0:
                raise Exception(f"Lỗi chuyển đổi Calibre: {result.stderr}")
                
        except FileNotFoundError:
            raise Exception("Calibre chưa được cài đặt hoặc không tìm thấy trong PATH. "
                          "Vui lòng cài đặt Calibre và thêm vào PATH để chuyển đổi sang MOBI.")

    def conversion_successful(self, output_path, output_format):
        """
        Hiển thị thông báo khi chuyển đổi thành công.
        """
        self.progress.stop()
        self.status_label.config(text=f"Chuyển đổi thành công! Đã lưu tại: {output_path}")
        messagebox.showinfo("Hoàn tất", f"File {output_format} đã được tạo thành công!\n\nĐường dẫn: {output_path}")
        self.convert_button.config(state=tk.NORMAL)

    def conversion_failed(self, error):
        """
        Hiển thị thông báo khi có lỗi xảy ra.
        """
        self.progress.stop()
        self.status_label.config(text=f"Lỗi: {error}")
        messagebox.showerror("Lỗi chuyển đổi", f"Đã có lỗi xảy ra:\n{error}")
        self.convert_button.config(state=tk.NORMAL)

if __name__ == "__main__":
    root = tk.Tk()
    app = EbookAllInOneConverterApp(root)
    root.mainloop()
